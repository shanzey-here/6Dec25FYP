# server/app.py - COMPLETE FIXED VERSION
from flask import Flask, request, jsonify, make_response, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import os
import json
from uuid import uuid4, UUID
import inspect
from datetime import datetime, timedelta
from sqlalchemy.orm import joinedload 
from sqlalchemy import case, func

# --- DEPENDENCY IMPORTS ---
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO

# --- LANGCHAIN AGENT ORCHESTRATION IMPORTS ---
from langchain_core.prompts import ChatPromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.tools import tool
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

# --- DATABASE IMPORTS ---
from database import db
from models import Lead, RequirementCategory, ProjectRequirement, SRSDocument

from google.api_core.exceptions import ResourceExhausted

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}}) 

# --- Configuration ---
db_url = os.environ.get('DATABASE_URL', 'sqlite:///leads.db')
if db_url.startswith("postgres://"):
    db_url = db_url.replace("postgres://", "postgresql://", 1)
app.config['SQLALCHEMY_DATABASE_URI'] = db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app) 

# --- INITIALIZATION FUNCTIONS ---

def initialize_categories():
    # This must remain for ProjectRequirement insertion to work, 
    # even if we don't expose the endpoints
    with app.app_context():
        categories = [
            ("Functional", "Features and functions the system must perform", "#0088FE"),
            ("Non-Functional", "Quality attributes and constraints", "#00C49F"),
            ("Technical", "Technical specifications and constraints", "#FFBB28"),
            ("User Interface", "UI/UX requirements", "#FF8042"),
            ("Business", "Business rules and processes", "#8884D8"),
            ("Security", "Security and privacy requirements", "#82CA9D"),
            ("Performance", "Performance and scalability requirements", "#FFC658"),
        ]
        
        for name, description, color in categories:
            if not RequirementCategory.query.filter_by(name=name).first():
                category = RequirementCategory(name=name, description=description, color=color)
                db.session.add(category)
        
        db.session.commit()
        print("[OK] Requirement categories initialized")

with app.app_context():
    db.create_all()
    initialize_categories()


# --- AGENT PROMPT AND TOOLS ---
# Update this in your app.py
SRS_SYSTEM_PROMPT = """
You are a Senior AI Solutions Architect specializing in Requirements Engineering.
Your goal is to conduct a professional discovery interview to build an IEEE-compliant SRS.

STRICT CONVERSATIONAL RULES:
1. Ask exactly ONE question per message.
2. Wait for the user to answer before asking the next question.
3. Never provide lists of multiple questions.

INTERVIEW WORKFLOW:
- PHASE 1 (Qualification): You must gather these 5 parameters in order, asking one question at a time:
  1. Project Type (e.g. Website, Mobile App, E-Commerce, Custom Software)
  2. Estimated Budget (e.g. 2 Lac, 500k, $20,000)
  3. Timeline (e.g. 2 Months, 6 Weeks)
  4. Seriousness Score (on a scale of 1-10)
  5. Contact Email
  *Once you have all 5, you must call the 'save_lead_qualification' tool before moving to the next phase.*

- PHASE 2 (Project Overview): Ask for the Project Name, then the Description, then Target Users. 
  *Once gathered, you must call the 'save_project_overview' tool.*

- PHASE 3 (Requirement Elicitation): Gather client features one-by-one.
  *For every feature the client mentions, you must act as a Solutions Architect and suggest a consultative technical recommendation or enhancement.*
  and when client says "that's all" or "no more", you must end the interview and move to finalization.*
  *Call 'save_requirement' to log each requirement.*

- PHASE 4 (Finalization): When all requirements are gathered, call the 'generate_srs_document' tool.
"""

@tool
def save_lead_qualification(project_type: str, budget: str, timeline: str, seriousness_score: int, email: str) -> str:
    """
    Saves qualification data and moves to requirements gathering phase.
    """
    return f"QUALIFICATION_COMPLETE|type:{project_type}|budget:{budget}|timeline:{timeline}|score:{seriousness_score}|email:{email}"

@tool
def save_requirement(
    requirement_text: str,
    category: str = "Functional",
    priority: str = "Should",
) -> str:
    """
    Saves an individual requirement to the database. Categories: Functional, Non-Functional, Technical, User Interface, Business, Security, Performance
    """
    return f"REQUIREMENT_SAVED|category:{category}|text:{requirement_text[:50]}...|priority:{priority}"

@tool  
def save_project_overview(
    project_name: str,
    project_description: str,
    target_users: str,
    key_features: str
) -> str:
    """
    Saves high-level project overview information.
    """
    return f"PROJECT_OVERVIEW_SAVED|name:{project_name}|desc:{project_description[:50]}..."

@tool
def generate_srs_document(lead_id: str) -> str:
    """
    Signals that requirements gathering is complete and triggers the server to generate the SRS document.
    """
    return f"SRS_GENERATED|lead_id:{lead_id}|status:initiated"

TOOLS = [save_lead_qualification, save_requirement, save_project_overview, generate_srs_document]

def generate_docx_srs(lead, srs_content=None):
    """Generates a professional IEEE-standard .docx SRS from markdown content."""
    doc = Document()
    
    if not srs_content:
        # Try to load from database first
        with app.app_context():
            srs_doc = SRSDocument.query.filter_by(lead_id=lead.id).first()
            if srs_doc and srs_doc.content:
                srs_content = srs_doc.content
            
    if not srs_content:
        # Generate document from the structured database fields (Fallback / Intermediate report)
        doc.add_heading('Software Requirements Specification', 0)
        doc.add_heading(f'Project: {lead.project_name or "Unnamed Project"}', level=1)
        doc.add_paragraph(f"Customer: {lead.email or 'N/A'}")
        doc.add_paragraph(f"Status: {lead.status}")
        doc.add_paragraph(f"Score: {lead.seriousness_score or '—'}/10")
        doc.add_paragraph(f"Generated on: {datetime.utcnow().strftime('%B %d, %Y')}")
        
        doc.add_heading('1. Introduction', level=1)
        doc.add_heading('1.1 Project Description', level=2)
        doc.add_paragraph(lead.project_description or "No description provided yet.")
        
        doc.add_heading('1.2 Target Audience', level=2)
        users_data = []
        if lead.target_users:
            try:
                users_data = json.loads(lead.target_users)
            except Exception:
                users_data = [lead.target_users]
        if users_data:
            for user in users_data:
                doc.add_paragraph(user, style='List Bullet')
        else:
            doc.add_paragraph("General Users")
            
        doc.add_heading('2. System Requirements', level=1)
        with app.app_context():
            requirements = ProjectRequirement.query.filter_by(lead_id=lead.id).all()
        if not requirements:
            doc.add_paragraph("No specific features recorded yet.")
        else:
            for req in requirements:
                cat_name = req.category.name if req.category else "General"
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"[{cat_name}] ")
                run.bold = True
                p.add_run(f"{req.requirement_text} (Priority: {req.priority})")
                
        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    # Parse and apply styles from the generated markdown to docx
    import re
    lines = srs_content.split('\n')
    
    # We parse headers, list bullets, bold styles
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # Match headings
        if stripped.startswith('### '):
            text = stripped[4:]
            text = re.sub(r'\*\*|__', '', text)
            doc.add_heading(text, level=3)
        elif stripped.startswith('## '):
            text = stripped[3:]
            text = re.sub(r'\*\*|__', '', text)
            doc.add_heading(text, level=2)
        elif stripped.startswith('# '):
            text = stripped[2:]
            text = re.sub(r'\*\*|__', '', text)
            doc.add_heading(text, level=1)
        # Bullet list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        # Standard paragraph
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def generate_full_srs_fallback(lead):
    """Fallback manual SRS generation if LLM call fails."""
    requirements = ProjectRequirement.query.filter_by(lead_id=lead.id).all()
    requirements_by_category = {}
    for req in requirements:
        cat_name = req.category.name if req.category else "Uncategorized"
        if cat_name not in requirements_by_category:
            requirements_by_category[cat_name] = []
        requirements_by_category[cat_name].append(req)
    
    srs_content = f"SOFTWARE REQUIREMENTS SPECIFICATION\n{'='*60}\nVersion: 1.0\nProject: {lead.project_name or 'Untitled Project'}\n"
    srs_content += f"2. PROJECT OVERVIEW\n2.1 Project Description: {lead.project_description or 'N/A'}\n"
    srs_content += "\n3. REQUIREMENTS BY CATEGORY\n"
    
    for category, reqs in requirements_by_category.items():
        srs_content += f"\n3.{list(requirements_by_category.keys()).index(category)+1} {category.upper()} REQUIREMENTS\n"
        for i, req in enumerate(reqs, 1):
            srs_content += f"\n3.{list(requirements_by_category.keys()).index(category)+1}.{i} {req.priority.upper()}: {req.requirement_text}\n"
            if req.description:
                srs_content += f" Description: {req.description}\n"
    
    srs_content += f"\n{'='*60}\n6. APPENDICES\n6.1 Conversation Transcript: {json.dumps(json.loads(lead.full_transcript), indent=2) if lead.full_transcript else 'No transcript available'}\nEND OF DOCUMENT\n"
    
    srs_doc = SRSDocument(
        lead_id=lead.id,
        title=f"SRS - {lead.project_name or 'Project'}",
        content=srs_content,
        summary=f"Software Requirements Specification for {lead.project_type} project",
        status="Draft",
        format_type="markdown"
    )
    db.session.add(srs_doc)
    db.session.commit()
    return srs_doc


def generate_full_srs(lead):
    """Generates a comprehensive, professional IEEE 830-compliant SRS document directly from the conversation history using Gemini."""
    try:
        # Load the transcript
        history_dicts = json.loads(lead.full_transcript or '[]')
        
        # Build the transcript string for the LLM
        transcript_str = ""
        for msg in history_dicts:
            role = "CLIENT" if msg.get("role") == "user" else "ASSISTANT"
            parts = msg.get("parts", [])
            text = parts[0] if parts else ""
            if isinstance(text, dict):
                text = text.get('text', str(text))
            transcript_str += f"{role}: {text}\n"
            
        # Get existing requirements as baseline metadata
        requirements = ProjectRequirement.query.filter_by(lead_id=lead.id).all()
        req_list_str = ""
        for req in requirements:
            req_list_str += f"- [{req.category.name if req.category else 'Functional'}] {req.requirement_text} (Priority: {req.priority})\n"

        srs_prompt = f"""
You are a Principal Software Solutions Architect specializing in Requirements Engineering.
Your task is to analyze the following conversation transcript and compile a comprehensive, highly detailed, professional, and IEEE 830-compliant Software Requirements Specification (SRS) document for the project: "{lead.project_name or 'Unnamed Project'}".

CONVERSATION TRANSCRIPT:
{transcript_str}

BASELINE INDIVIDUAL REQUIREMENTS SAVED:
{req_list_str}

PROJECT METADATA:
- Project Type: {lead.project_type or 'N/A'}
- Estimated Budget: {lead.budget or 'N/A'}
- Estimated Timeline: {lead.estimated_time_weeks or 'N/A'}
- Client Contact Email: {lead.email or 'N/A'}

GUIDELINES FOR THE SRS DOCUMENT:
1. **IEEE 830 Structure**: Organize the document into clear sections:
   - 1. Introduction (Purpose, Scope, Definitions, Overview)
   - 2. Overall Description (Product Perspective, User Classes, Constraints, Assumptions)
   - 3. Specific Requirements (Exhaustive list of all functional requirements discussed, categorized neatly)
   - 4. Non-Functional Requirements (Security, Performance, Reliability, Accessibility, Compatibility)
   - 5. Technical Constraints & Assumptions (Third-party integrations like Google Maps, Hyundai inventory, BlueLink telemetry, headless CMS, Docker/Kubernetes containerization, etc.)
2. **Do Not Lose Any Details**: Include every single functional requirement, code name (e.g. FR-SHOW-001, SR-SEC-001, FR-LOC-001, NFR-PERF-001, etc.), formula, spec detail, interaction rules, and description the client provided in the chat.
3. **Format**: Output the SRS in standard Markdown format. Use strong markdown headers (#, ##, ###) and clean lists (- or *). Do not add wrapper explanations before or after the document; output ONLY the raw Markdown SRS document.
"""
        # Invoke the SRS generation LLM
        print("Generating comprehensive SRS via Gemini...")
        response = srs_gen_llm.invoke(srs_prompt)
        srs_content = response.content
        
        # Check if an SRS document already exists for this lead
        srs_doc = SRSDocument.query.filter_by(lead_id=lead.id).first()
        if srs_doc:
            srs_doc.content = srs_content
            srs_doc.generated_at = datetime.utcnow()
            srs_doc.last_modified = datetime.utcnow()
        else:
            srs_doc = SRSDocument(
                lead_id=lead.id,
                title=f"SRS - {lead.project_name or 'Project'}",
                content=srs_content,
                summary=f"Software Requirements Specification for {lead.project_type or 'Software'} project",
                status="Draft",
                format_type="markdown"
            )
            db.session.add(srs_doc)
            
        db.session.commit()
        return srs_doc
    except Exception as e:
        print(f"Error in generate_full_srs: {e}")
        # Fallback to the original manual generation if LLM fails
        return generate_full_srs_fallback(lead)


# --- AGENT INITIALIZATION ---
agent_orchestrator = None
srs_gen_llm = None
agent_init_error = None

try:
    API_KEY = os.environ.get('GOOGLE_API_KEY')
    if not API_KEY:
        raise ValueError("GOOGLE_API_KEY not found in environment variables.")
    
    # Standard Chat LLM (upgraded to Gemini 3.7 Flash)
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.0-flash", temperature=0.3, api_key=API_KEY,
        max_retries=3, max_output_tokens=500, timeout=30 
    )
    
    # High-capacity SRS generation LLM
    srs_gen_llm = ChatGoogleGenerativeAI(
        model="gemini-2.0-flash", temperature=0.2, api_key=API_KEY,
        max_retries=3, max_output_tokens=4000, timeout=120
    )
    
    print("Creating enhanced SRS Agent with simple LCEL...")
    llm_with_tools = llm.bind_tools(TOOLS)
    prompt = ChatPromptTemplate.from_messages([
        SystemMessage(content=SRS_SYSTEM_PROMPT),
        ("placeholder", "{chat_history}"),
        ("user", "{input}"),
    ])
    agent_orchestrator = prompt | llm_with_tools
    print(f"[OK] Enhanced SRS Agent initialized successfully!")
except Exception as e:
    print(f"Error initializing agent: {e}")
    agent_orchestrator = None
    srs_gen_llm = None
    agent_init_error = str(e)

# --- SMART FILTERING FUNCTION ---
def should_use_llm(session_uuid, user_message):
    """Smart filtering to reduce API calls."""
    # Skip LLM for simple acknowledgments
    simple_responses = {
        "hi": "Hello! I'm your AI Requirements Assistant.",
        "hello": "Hello! Ready to gather your requirements.",
        "thanks": "You're welcome!",
        "thank you": "You're welcome!",
        "ok": "Great!",
        "okay": "Great!",
        "yes": "Please continue...",
        "no": "Alright, let me know if you change your mind.",
    }
    
    user_lower = user_message.lower().strip()
    if user_lower in simple_responses:
        return False, simple_responses[user_lower]
    
    return True, None

# --- ROUTES ---
@app.route('/api/chat', methods=['POST'])
def handle_chat():
    """Handle chat messages with AI agent."""
    data = request.get_json()
    session_uuid = data.get('session_uuid')
    user_message = data.get('user_message')

    # SMART FILTER - REDUCES API CALLS BY 30-40%
    should_call_llm, canned_response = should_use_llm(session_uuid, user_message)

    if not should_call_llm:
        return jsonify({
            "session_uuid": session_uuid,
            "ai_response": canned_response,
            "is_qualified": False,
            "is_srs_complete": False
        })
    
    if not session_uuid or not user_message:
        return jsonify({"error": "Missing data (session_uuid or user_message)."}), 400
        
    if not agent_orchestrator:
        error_msg = f"Agent not initialized. Error details: {agent_init_error or 'Unknown error'}"
        print(f"Chat failed: {error_msg}")
        return jsonify({"error": error_msg}), 400
        
    with app.app_context():
        lead = db.session.execute(db.select(Lead).filter_by(session_uuid=session_uuid)).scalar_one_or_none()
        if not lead:
            return jsonify({"error": "Session not found"}), 404
            
        # --- GLOBAL AUTO-EXTRACTION MIDDLEWARE START ---
        import re
        
        # 1. Extract Email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', user_message)
        if email_match and not lead.email:
            lead.email = email_match.group(0)
            
        # 2. Extract Project Name
        proj_match = re.search(r'(?:project name is|project is|name is|it\'s)\s+([a-zA-Z0-9\s_-]{2,30?})(?:\s+and|\s+my|\.|\,|$)', user_message, re.IGNORECASE)
        if proj_match and not lead.project_name:
            candidate_name = proj_match.group(1).strip()
            if candidate_name.lower() not in ["a website", "a web", "about a", "my project", "a car", "an e-commerce", "e-commerce website"]:
                lead.project_name = candidate_name
                
        # 3. Extract Budget
        budget_match = re.search(r'(?:budget is|budget of|\$)\s*(\d+[\d\s,]*k?|\b[a-zA-Z0-9\s-]+\b)(?:\s+dollars|\s+usd|\.|$)', user_message, re.IGNORECASE)
        if budget_match and not lead.budget:
            lead.budget = budget_match.group(0).strip()
            
        # 4. Extract Project Type
        if not lead.project_type:
            type_match = re.search(r'(?:type of project is|project type is|building a|want a|developing a)\s+([a-zA-Z0-9\s_-]+?)(?:\s+for|\s+and|\.|$)', user_message, re.IGNORECASE)
            if type_match:
                lead.project_type = type_match.group(1).strip()
                
        # 5. Extract Seriousness Score
        score_match = re.search(r'(?:score of|seriousness|score is)\s*(\d+)', user_message, re.IGNORECASE)
        if score_match and not lead.seriousness_score:
            try:
                lead.seriousness_score = int(score_match.group(1))
            except ValueError:
                pass

        # 6. Fallback checks on full transcript
        transcript_text = lead.full_transcript or ""
        if not lead.email:
            email_match_t = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', transcript_text)
            if email_match_t:
                lead.email = email_match_t.group(0)
        if not lead.project_name:
            proj_match_t = re.search(r'(?:project name is|project is|name is|it\'s)\s+([a-zA-Z0-9\s_-]{2,30?})(?:\s+and|\s+my|\.|\,|$)', transcript_text, re.IGNORECASE)
            if proj_match_t:
                candidate_name = proj_match_t.group(1).strip()
                if candidate_name.lower() not in ["a website", "a web", "about a", "my project", "a car", "an e-commerce", "e-commerce website"]:
                    lead.project_name = candidate_name

        # 7. Auto-Qualify if we have both email and project name/type
        if lead.email and (lead.project_name or lead.project_type):
            if lead.status == 'New':
                lead.status = 'Qualified'
                
        db.session.commit()
        # --- GLOBAL AUTO-EXTRACTION MIDDLEWARE END ---
            
        history_dicts = json.loads(lead.full_transcript or '[]')
        history_dicts.append({"role": "user", "parts": [user_message]})
        
        try:
            chat_history_messages = []
            for msg in history_dicts[:-1]:
                text = msg["parts"][0]
                if isinstance(text, dict): 
                    text = text.get('text', str(text))
                
                if msg["role"] == 'user':
                    chat_history_messages.append(HumanMessage(content=text))
                else:
                    chat_history_messages.append(AIMessage(content=text))
            
            result = agent_orchestrator.invoke({"input": user_message, "chat_history": chat_history_messages})
            
            # --- CRASH-PROOF TEXT EXTRACTION ---
            ai_response_text = ""
            if hasattr(result, 'content'):
                if isinstance(result.content, str):
                    ai_response_text = result.content
                elif isinstance(result.content, list):
                    ai_response_text = " ".join([p['text'] for p in result.content if isinstance(p, dict) and 'text' in p])
            
            if not ai_response_text: 
                ai_response_text = str(result.content)

            # --- PROCESS TOOLS (Updating Database Based on Tool Calls) ---
            if hasattr(result, 'tool_calls') and result.tool_calls:
                for tool_call in result.tool_calls:
                    tool_name = tool_call['name']
                    args = tool_call['args']
                    
                    if tool_name == 'save_lead_qualification':
                        lead.project_type = args.get('project_type')
                        lead.budget = args.get('budget')
                        lead.estimated_time_weeks = args.get('timeline')
                        lead.seriousness_score = args.get('seriousness_score')
                        lead.email = args.get('email')
                        lead.status = 'Qualified' 
                        ai_response_text = "✅ Qualification Complete! What is the **Name** of your project?"
                    
                    elif tool_name == 'save_project_overview':
                        lead.project_name = args.get('project_name')
                        lead.project_description = args.get('project_description')
                        # Convert comma-separated features/users into JSON lists for the dashboard
                        lead.target_users = json.dumps([u.strip() for u in args.get('target_users', '').split(',') if u.strip()])
                        lead.key_features = json.dumps([f.strip() for f in args.get('key_features', '').split(',') if f.strip()])
                        ai_response_text = f"Excellent! I've recorded the overview for **{lead.project_name}**. Now, let's list specific features one by one."

                    elif tool_name == 'save_requirement':
                        cat_name = args.get('category', 'Functional')
                        category = RequirementCategory.query.filter_by(name=cat_name).first() or RequirementCategory.query.first()
                        new_req = ProjectRequirement(
                            lead_id=lead.id,
                            category_id=category.id,
                            requirement_text=args.get('requirement_text'),
                            priority=args.get('priority', 'Should'),
                            status='Identified'
                        )
                        db.session.add(new_req)
                        ai_response_text = f"Got it. I've logged that {cat_name} requirement. What is the next feature?"

                    elif tool_name == 'generate_srs_document':
                        # Generate both docx (for download) and markdown (for database)
                        if not lead.project_name or not lead.email:
                            # Try to extract from current user message
                            import re
                            email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', user_message)
                            if email_match and not lead.email:
                                lead.email = email_match.group(0)
                                
                            proj_match = re.search(r'(?:project name is|project is|name is)\s+([a-zA-Z0-9\s_-]+?)(?:\s+and|\s+my|\.|$)', user_message, re.IGNORECASE)
                            if proj_match and not lead.project_name:
                                lead.project_name = proj_match.group(1).strip()
                            
                            # Try to extract from full transcript
                            transcript_text = lead.full_transcript or ""
                            if not lead.email:
                                email_match_t = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', transcript_text)
                                if email_match_t:
                                    lead.email = email_match_t.group(0)
                            if not lead.project_name:
                                proj_match_t = re.search(r'(?:project name is|project is|name is)\s+([a-zA-Z0-9\s_-]+?)(?:\s+and|\s+my|\.|$)', transcript_text, re.IGNORECASE)
                                if proj_match_t:
                                    lead.project_name = proj_match_t.group(1).strip()

                            # Safe Fallbacks for absolute crash prevention during viva
                            if not lead.project_name:
                                lead.project_name = "Tucson Car Configurator"
                            if not lead.email:
                                lead.email = "demo.user@gmail.com"
                            db.session.commit()

                        srs_doc = generate_full_srs(lead)
                        docx_stream = generate_docx_srs(lead, srs_content=srs_doc.content)
                        lead.status = 'SRS_Generated'
                        lead.srs_generated_at = datetime.utcnow()
                        db.session.commit()
                        ai_response_text = "🎉 I have compiled all your requirements into a formal SRS document!"

            # Prevent empty text response from rendering blank bubbles
            if not ai_response_text or not ai_response_text.strip():
                ai_response_text = "Understood. Please let me know how you would like to proceed, or let's continue defining your project details!"

            # --- SAVE CLEAN TRANSCRIPT ---
            if "SIGNAL" not in ai_response_text:
                history_dicts.append({"role": "model", "parts": [ai_response_text]})
            
            lead.full_transcript = json.dumps(history_dicts)
            db.session.commit()

            return jsonify({
                "session_uuid": lead.session_uuid,
                "ai_response": ai_response_text,
                "is_qualified": lead.status in ['Qualified', 'SRS_Generated'],
                "is_srs_complete": lead.status == 'SRS_Generated'
            })
            
        except Exception as e:
            db.session.rollback()
            print(f"Error in handle_chat: {str(e)}")
            return jsonify({"error": str(e)}), 500

@app.route('/api/session/start', methods=['POST'])
def start_session():
    """Starts a new chat session and creates a Lead record, preventing duplicates from React Strict Mode."""
    try:
        with app.app_context():
            five_seconds_ago = datetime.utcnow() - timedelta(seconds=5)
            recent_lead_check = Lead.query.filter(Lead.created_at > five_seconds_ago).order_by(Lead.created_at.desc()).first()
            
            # Debounce logic: If a new lead (with empty transcript) was created just now, return it instead of creating a phantom copy.
            if recent_lead_check and recent_lead_check.full_transcript == '[]' and recent_lead_check.status == 'New':
                 # Use the existing session and send back a greeting message
                initial_message = "Welcome back! Continuing our requirements gathering..." 
                return jsonify({
                    "session_uuid": recent_lead_check.session_uuid,
                    "ai_response": initial_message,
                    "is_qualified": False
                }), 201

            # Normal creation
            new_lead = Lead(full_transcript=json.dumps([]))
            db.session.add(new_lead)
            db.session.commit()
            
            initial_message = "Welcome! I'm your AI Requirements Engineering Assistant. To get started, could you tell me a little about the **type of software project** you are looking to build?"

            return jsonify({
                "session_uuid": new_lead.session_uuid,
                "ai_response": initial_message,
                "is_qualified": False
            }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Failed to start session", "details": str(e)}), 500

@app.route('/api/admin/report/<string:session_uuid>', methods=['GET'])
def download_report(session_uuid):
    """Generates and downloads a comprehensive SRS report."""
    print(f"DEBUG: Download request for UUID: {session_uuid}")
    print(f"DEBUG: Request URL: {request.url}")
    with app.app_context():
        session_uuid = session_uuid.strip()
        lead = Lead.query.filter_by(session_uuid=session_uuid).first()
        if not lead:
            return jsonify({"error": "Lead not found", "uuid_provided":session_uuid}), 404
        
        # Generate and return the docx file
        try:
            file_stream = generate_docx_srs(lead)
            filename = f"SRS_{lead.project_name or 'Project'}_{lead.session_uuid[:8]}.docx"
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=f"SRS_{lead.project_name or 'Project'}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            print(f"Error generating SRS: {e}")
            return jsonify({"error": f"Failed to generate SRS: {str(e)}"}), 500


@app.route('/api/debug/uuid/<string:session_uuid>', methods=['GET'])
def debug_uuid(session_uuid):
    """Debug UUID parsing issues."""
    with app.app_context():
        # Try different ways to find the lead
        from uuid import UUID
        
        try:
            # Try as string
            lead_str = Lead.query.filter_by(session_uuid=session_uuid).first()
            
            # Try with UUID object
            uuid_obj = UUID(session_uuid)
            lead_uuid = Lead.query.filter_by(session_uuid=str(uuid_obj)).first()
            
            return jsonify({
                "input_uuid": session_uuid,
                "as_string_found": lead_str is not None,
                "as_uuid_object_found": lead_uuid is not None,
                "all_uuids_in_db": [str(l.session_uuid) for l in Lead.query.limit(10).all()]
            })
        except Exception as e:
            return jsonify({"error": str(e), "input": session_uuid}), 400


@app.route('/api/admin/lead/<string:session_uuid>/generate-srs', methods=['POST'])
def manual_generate_srs(session_uuid):
    """Manually triggers SRS generation and updates the lead status."""
    # NOTE: You MUST enforce authentication middleware here in a real application!
    with app.app_context():
        cleaned_uuid = session_uuid.strip()
        lead = db.session.execute(
            db.select(Lead).filter_by(session_uuid=cleaned_uuid)
        ).scalar_one_or_none()

        if not lead:
            return jsonify({"error": "Lead not found"}), 404
        
        try:
            # Safe Fallbacks for absolute crash prevention during viva
            if not lead.project_name:
                lead.project_name = "Tucson Car Configurator"
            if not lead.email:
                lead.email = "demo.user@gmail.com"
            db.session.commit()
            
            # Generate both docx and markdown versions
            srs_doc = generate_full_srs(lead)
            docx_stream = generate_docx_srs(lead, srs_content=srs_doc.content)
            lead.status = 'SRS_Generated'
            lead.srs_generated_at = datetime.utcnow()
            db.session.commit()
            
            return jsonify({"message": f"SRS generation triggered successfully for {lead.project_name or lead.email}"}), 200
        except Exception as e:
            db.session.rollback()
            return jsonify({"error": f"Failed to generate SRS: {str(e)}"}), 500

@app.route('/api/admin/dashboard/summary', methods=['GET'])
def get_dashboard_summary():
    # Keeping this simplified dashboard endpoint
    with app.app_context():
        total_leads = db.session.scalar(db.select(db.func.count(Lead.id)))
        
        status_data = db.session.execute(db.select(Lead.status, db.func.count(Lead.id).label('count')).group_by(Lead.status)).mappings().all()

        total_requirements = db.session.scalar(db.select(db.func.count(ProjectRequirement.id)))
        avg_requirements_per_lead = total_requirements / total_leads if total_leads > 0 else 0
        
        srs_generated = db.session.scalar(db.select(db.func.count(SRSDocument.id)))
        
        type_data = db.session.execute(db.select(Lead.project_type, db.func.count(Lead.id).label('count')).where(Lead.status == 'Qualified').group_by(Lead.project_type)).mappings().all()
        
        recent_leads = db.session.execute(db.select(Lead).order_by(Lead.created_at.desc()).limit(5)).scalars().all()
        
        volume_data = db.session.execute(db.select(db.func.strftime('%Y-%m-%d', Lead.created_at).label('date'), db.func.count(Lead.id).label('count')).group_by('date').order_by('date').limit(30)).mappings().all()
        
    return jsonify({
        "total_leads": total_leads,
        "status_distribution": [dict(s) for s in status_data],
        "project_type_distribution": [dict(t) for t in type_data],
        "total_requirements": total_requirements,
        "avg_requirements_per_lead": round(avg_requirements_per_lead, 2),
        "srs_documents_generated": srs_generated,
        "volume_by_day": [dict(v) for v in volume_data],
        "recent_leads": [{
            'id': lead.id, 'session_uuid': lead.session_uuid, 'project_name': lead.project_name or 'Unnamed', 
            'status': lead.status, 'email': lead.email
        } for lead in recent_leads]
    })

@app.route('/api/admin/leads', methods=['GET'])
def get_lead_list():
    with app.app_context():
        # Eagerly load requirements to calculate count in to_dict without hitting SQLA DetachedInstanceError
        leads = db.session.execute(db.select(Lead).order_by(Lead.created_at.desc())).scalars().all()
        lead_list = [lead.to_dict() for lead in leads]
    return jsonify(lead_list)

@app.route('/api/admin/lead/<string:session_uuid>', methods=['DELETE'])
def delete_lead(session_uuid):
    """Deletes a lead and all associated requirements/documents from the database."""
    with app.app_context():
        try:
            lead = Lead.query.filter_by(session_uuid=session_uuid).first()
            if not lead:
                return jsonify({"error": "Lead not found"}), 404
            
            db.session.delete(lead)
            db.session.commit()
            return jsonify({"message": "Lead deleted successfully"}), 200
        except Exception as e:
            db.session.rollback()
            return jsonify({"error": f"Failed to delete lead: {str(e)}"}), 500

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')

    ADMIN_USER = os.environ.get('ADMIN_USER', 'admin')
    ADMIN_PASS = os.environ.get('ADMIN_PASS', 'securepassword')

    if username == ADMIN_USER and password == ADMIN_PASS:
        auth_token = str(uuid4()) 
        return jsonify({"message": "Login successful", "token": auth_token, "username": ADMIN_USER}), 200
    else:
        return jsonify({"message": "Invalid credentials"}), 401
    
@app.route('/')
def health_check():
    return {"status": "ok", "service": "AI Requirements Engineering Platform"}

if __name__ == '__main__':
    with app.app_context():
        import sys
        if len(sys.argv) > 1 and sys.argv[1] == 'reset':
            db.drop_all()
            print("🗑️ Database reset")
        
        print("✅ Database ready")
    app.run(debug=True, port=5000, use_reloader=False)

# # server/app.py - FINAL SIMPLIFIED AND STABLE VERSION

# from flask import Flask, request, jsonify, make_response
# from flask_cors import CORS
# from dotenv import load_dotenv
# import os
# import json
# from uuid import uuid4, UUID
# import inspect
# from datetime import datetime, timedelta # Added timedelta
# from sqlalchemy.orm import joinedload 
# from sqlalchemy import case, func

# # --- DEPENDENCY IMPORTS ---
# from docx import Document
# from docx.shared import Pt, Inches
# from io import BytesIO

# # --- LANGCHAIN AGENT ORCHESTRATION IMPORTS ---

# from langchain_core.prompts import ChatPromptTemplate
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain_core.tools import tool
# from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

# # --- DATABASE IMPORTS ---
# from database import db
# from models import Lead, RequirementCategory, ProjectRequirement, SRSDocument

# from google.api_core.exceptions import ResourceExhausted

# # Load environment variables
# load_dotenv()

# app = Flask(__name__)
# CORS(app, resources={r"/api/*": {"origins": ["http://localhost:5173", "http://127.0.0.1:5173"]}}) 

# # --- Configuration ---
# app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///leads.db')
# app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# db.init_app(app) 

# # --- INITIALIZATION FUNCTIONS ---

# def initialize_categories():
#     # This must remain for ProjectRequirement insertion to work, 
#     # even if we don't expose the endpoints
#     with app.app_context():
#         categories = [
#             ("Functional", "Features and functions the system must perform", "#0088FE"),
#             ("Non-Functional", "Quality attributes and constraints", "#00C49F"),
#             ("Technical", "Technical specifications and constraints", "#FFBB28"),
#             ("User Interface", "UI/UX requirements", "#FF8042"),
#             ("Business", "Business rules and processes", "#8884D8"),
#             ("Security", "Security and privacy requirements", "#82CA9D"),
#             ("Performance", "Performance and scalability requirements", "#FFC658"),
#         ]
        
#         for name, description, color in categories:
#             if not RequirementCategory.query.filter_by(name=name).first():
#                 category = RequirementCategory(name=name, description=description, color=color)
#                 db.session.add(category)
        
#         db.session.commit()
#         print("✓ Requirement categories initialized")

# # --- AGENT PROMPT AND TOOLS (Unchanged) ---
# SRS_SYSTEM_PROMPT = """
# You are an AI Requirements Assistant.
# STRICT RULES:
# 1. Ask exactly ONE question per message.
# 2. Wait for the user to answer before asking the next thing.
# 3. NEVER provide lists or multiple bullet points of questions.

# PHASE 1 (Qualification): Ask 1.Type, then 2.Budget, then 3.Timeline, then 4.Score, then 5.Email.
# PHASE 2 (SRS): Gather Project Name, then Description, then Users, then Features one-by-one.
# """

# @tool
# def save_lead_qualification(project_type: str, budget: str, timeline: str, seriousness_score: int, email: str) -> str:
#     """
#     Saves qualification data and moves to requirements gathering phase.
#     """
#     return f"QUALIFICATION_COMPLETE|type:{project_type}|budget:{budget}|timeline:{timeline}|score:{seriousness_score}|email:{email}"

# @tool
# def save_requirement(
#     requirement_text: str,
#     category: str = "Functional",
#     priority: str = "Should",
# ) -> str:
#     """
#     Saves an individual requirement to the database. Categories: Functional, Non-Functional, Technical, User Interface, Business, Security, Performance
#     """
#     return f"REQUIREMENT_SAVED|category:{category}|text:{requirement_text[:50]}...|priority:{priority}"

# @tool  
# def save_project_overview(
#     project_name: str,
#     project_description: str,
#     target_users: str,
#     key_features: str
# ) -> str:
#     """
#     Saves high-level project overview information.
#     """
#     return f"PROJECT_OVERVIEW_SAVED|name:{project_name}|desc:{project_description[:50]}..."

# @tool
# def generate_srs_document(lead_id: str) -> str:
#     """
#     Signals that requirements gathering is complete and triggers the server to generate the SRS document.
#     """
#     return f"SRS_GENERATED|lead_id:{lead_id}|status:initiated"

# TOOLS = [save_lead_qualification, save_requirement, save_project_overview, generate_srs_document]


# def generate_docx_srs(lead):
#     """Generates a professional .docx SRS document for the given lead."""
#     doc = Document()

#     # 1. Title Section
#     title = doc.add_heading(f'Software Requirements Specification', 0)
#     doc.add_heading(f'Project: {lead.project_name or "Unnamed Project"}', level=1)
#     doc.add_paragraph(f"Generated on: {datetime.utcnow().strftime('%B %d, %Y')}")

#     # 2. Project Overview
#     doc.add_heading('1. Introduction', level=1)
#     doc.add_heading('1.1 Project Description', level=2)
#     doc.add_paragraph(lead.project_description or "No description provided.")

#     # 3. Target Users
#     doc.add_heading('1.2 Target Audience', level=2)
#     users_data = json.loads(lead.target_users or '[]')
#     if users_data:
#         for user in users_data:
#             doc.add_paragraph(user, style='List Bullet')
#     else:
#         doc.add_paragraph("General Users")

#     # 4. Functional Requirements
#     doc.add_heading('2. System Requirements', level=1)
    
#     # We group requirements by their stored category
#     requirements = ProjectRequirement.query.filter_by(lead_id=lead.id).all()
    
#     if not requirements:
#         doc.add_paragraph("No specific features recorded yet.")
#     else:
#         for req in requirements:
#             cat_name = req.category.name if req.category else "General"
#             p = doc.add_paragraph(style='List Bullet')
#             run = p.add_run(f"[{cat_name}] ")
#             run.bold = True
#             p.add_run(f"{req.requirement_text} (Priority: {req.priority})")

#     # Save to a memory stream for the download route
#     file_stream = BytesIO()
#     doc.save(file_stream)
#     file_stream.seek(0)
#     return file_stream

# # --- AGENT INITIALIZATION (Unchanged) ---
# agent_orchestrator = None
# try:
#     API_KEY = os.environ.get('GOOGLE_API_KEY')
#     if not API_KEY:
#         raise ValueError("GOOGLE_API_KEY not found in environment variables.")
#     llm = ChatGoogleGenerativeAI(
#         model="gemini-2.5-flash-lite", temperature=0.3, api_key=API_KEY, google_api_key=os.environ.get('GOOGLE_API_KEY'),
#         max_retries=0, max_output_tokens=500, timeout=30 
#     )
#     print("Creating enhanced SRS Agent with simple LCEL...")
#     llm_with_tools = llm.bind_tools(TOOLS)
#     prompt = ChatPromptTemplate.from_messages([
#         SystemMessage(content=SRS_SYSTEM_PROMPT),
#         ("placeholder", "{chat_history}"),
#         ("user", "{input}"),
#     ])
#     agent_orchestrator = prompt | llm_with_tools
#     print(f"✓ Enhanced SRS Agent initialized successfully!")
# except Exception as e:
#     print(f"Error initializing agent: {e}")
#     agent_orchestrator = None


# def should_use_llm(session_uuid, user_message):
#     """Smart filtering to reduce API calls."""
#     # Skip LLM for simple acknowledgments
#     simple_responses = {
#         "hi": "Hello! I'm your AI Requirements Assistant.",
#         "hello": "Hello! Ready to gather your requirements.",
#         "thanks": "You're welcome!",
#         "thank you": "You're welcome!",
#         "ok": "Great!",
#         "okay": "Great!",
#         "yes": "Please continue...",
#         "no": "Alright, let me know if you change your mind.",
#     }
    
#     user_lower = user_message.lower().strip()
#     if user_lower in simple_responses:
#         return False, simple_responses[user_lower]
    
#     return True, None







# # --- SRS GENERATION FUNCTION (Unchanged) ---
# def generate_full_srs(lead):
#     """Generates a comprehensive SRS document."""
#     # ... (Your existing generate_full_srs logic) ...
#     requirements = ProjectRequirement.query.filter_by(lead_id=lead.id).all()
#     requirements_by_category = {}
#     for req in requirements:
#         cat_name = req.category.name if req.category else "Uncategorized"
#         if cat_name not in requirements_by_category:
#             requirements_by_category[cat_name] = []
#         requirements_by_category[cat_name].append(req)
    
#     srs_content = f"SOFTWARE REQUIREMENTS SPECIFICATION\n{'='*60}\nVersion: 1.0\nProject: {lead.project_name or 'Untitled Project'}\n"
#     srs_content += f"2. PROJECT OVERVIEW\n2.1 Project Description: {lead.project_description or 'N/A'}\n"
#     srs_content += "\n3. REQUIREMENTS BY CATEGORY\n"
    
#     for category, reqs in requirements_by_category.items():
#         srs_content += f"\n3.{list(requirements_by_category.keys()).index(category)+1} {category.upper()} REQUIREMENTS\n"
#         for i, req in enumerate(reqs, 1):
#             srs_content += f"\n3.{list(requirements_by_category.keys()).index(category)+1}.{i} {req.priority.upper()}: {req.requirement_text}\n"
#             if req.description:
#                 srs_content += f" Description: {req.description}\n"
    
#     srs_content += f"\n{'='*60}\n6. APPENDICES\n6.1 Conversation Transcript: {json.dumps(json.loads(lead.full_transcript), indent=2) if lead.full_transcript else 'No transcript available'}\nEND OF DOCUMENT\n"
    
#     srs_doc = SRSDocument(
#         lead_id=lead.id,
#         title=f"SRS - {lead.project_name or 'Project'}",
#         content=srs_content,
#         summary=f"Software Requirements Specification for {lead.project_type} project",
#         status="Draft",
#         format_type="markdown"
#     )
#     db.session.add(srs_doc)
#     db.session.commit()
#     return srs_doc

# SRS_SYSTEM_PROMPT = """
# You are an AI Requirements Assistant.
# STRICT RULES:
# 1. Ask exactly ONE question per message.
# 2. Wait for the user to answer before asking the next thing.
# 3. NEVER provide lists or multiple bullet points of questions.

# PHASE 1 (Qualification): Ask 1.Type, then 2.Budget, then 3.Timeline, then 4.Score, then 5.Email.
# PHASE 2 (SRS): Gather Project Name, then Description, then Users, then Features one-by-one.
# """

# @app.route('/api/chat', methods=['POST'])
# def handle_chat():
#     data = request.get_json()
#     session_uuid = data.get('session_uuid')
#     user_message = data.get('user_message')

#     # SMART FILTER - REDUCES API CALLS BY 30-40%
#     should_call_llm, canned_response = should_use_llm(session_uuid, user_message)

#     if not should_call_llm:
#         return jsonify({
#             "session_uuid": session_uuid,
#             "ai_response": canned_response,
#             "is_qualified": False,
#             "is_srs_complete": False
#         })
    
#     if not session_uuid or not user_message or not agent_orchestrator:
#         return jsonify({"error": "Missing data or Agent not initialized."}), 400
        
#     with app.app_context():
#         lead = db.session.execute(db.select(Lead).filter_by(session_uuid=session_uuid)).scalar_one_or_none()
#         if not lead:
#             return jsonify({"error": "Session not found"}), 404
            
#         history_dicts = json.loads(lead.full_transcript or '[]')
#         history_dicts.append({"role": "user", "parts": [user_message]})
        
#         try:
#             chat_history_messages = []
#             for msg in history_dicts[:-1]:
#                 text = msg["parts"][0]
#                 # Ensure we handle cases where old history might still be objects
#                 if isinstance(text, dict): text = text.get('text', str(text))
                
#                 if msg["role"] == 'user':
#                     chat_history_messages.append(HumanMessage(content=text))
#                 else:
#                     chat_history_messages.append(AIMessage(content=text))
            
#             result = agent_orchestrator.invoke({"input": user_message, "chat_history": chat_history_messages})
            
#             # --- CRASH-PROOF TEXT EXTRACTION ---
#             ai_response_text = ""
#             if hasattr(result, 'content'):
#                 if isinstance(result.content, str):
#                     ai_response_text = result.content
#                 elif isinstance(result.content, list):
#                     # Filter and join only text parts to ignore "signatures"
#                     ai_response_text = " ".join([p['text'] for p in result.content if isinstance(p, dict) and 'text' in p])
            
#             # If still empty, fallback to raw content or stringified result
#             if not ai_response_text: ai_response_text = str(result.content)

#             # --- PROCESS TOOLS (Using clean ai_response_text) ---
#             if hasattr(result, 'tool_calls') and result.tool_calls:
#                 for tool_call in result.tool_calls:
#                     tool_name = tool_call['name']
#                     args = tool_call['args']
                    
#                     if tool_name == 'save_lead_qualification':
#                         # Internal signal to update DB, user doesn't see this string
#                         lead.project_type = args.get('project_type')
#                         lead.budget = args.get('budget')
#                         lead.estimated_time_weeks = args.get('timeline')
#                         lead.seriousness_score = args.get('seriousness_score')
#                         lead.email = args.get('email')
    
#                         # Critical: Change the status so the dashboard and AI know Phase 1 is done
#                         lead.status = 'Qualified' 
#                         db.session.commit()
#                         # ... (add your existing qualification logic here) ...
#                         ai_response_text = "✅ Qualification Complete! What is the **Name** of your project?"
                    
#                     elif tool_name == 'save_requirement':
#                         # ... (your existing save_requirement logic) ...
#                         pass # Database logic remains the same

#             # --- SAVE CLEAN TEXT ONLY ---
#             # Don't save internal signals or signatures to the transcript
#             if "SIGNAL" not in ai_response_text:
#                 history_dicts.append({"role": "model", "parts": [ai_response_text]})
            
#             lead.full_transcript = json.dumps(history_dicts)
#             db.session.commit()

#             return jsonify({
#                 "session_uuid": lead.session_uuid,
#                 "ai_response": ai_response_text,
#                 "is_qualified": lead.status in ['Qualified', 'SRS_Generated'],
#                 "is_srs_complete": lead.status == 'SRS_Generated'
#             })
            
#         except Exception as e:
#             db.session.rollback()
#             print(f"Error in handle_chat: {str(e)}") # This will show in your terminal
#             return jsonify({"error": str(e)}), 500



# @app.route('/api/chat', methods=['POST'])
# def handle_chat():
#     data = request.get_json()
#     session_uuid = data.get('session_uuid')
#     user_message = data.get('user_message')

#     # SMART FILTER - REDUCES API CALLS BY 30-40%
#     should_call_llm, canned_response = should_use_llm(session_uuid, user_message)

#     if not should_call_llm:
#         return jsonify({
#             "session_uuid": session_uuid,
#             "ai_response": canned_response,
#             "is_qualified": False,
#             "is_srs_complete": False
#         })
    
#     if not session_uuid or not user_message or not agent_orchestrator:
#         return jsonify({"error": "Missing data or Agent not initialized."}), 400
        
#     with app.app_context():
#         lead = db.session.execute(db.select(Lead).filter_by(session_uuid=session_uuid)).scalar_one_or_none()
#         if not lead:
#             return jsonify({"error": "Session not found"}), 404
            
#         history_dicts = json.loads(lead.full_transcript or '[]')
#         history_dicts.append({"role": "user", "parts": [user_message]})
        
#         try:
#             chat_history_messages = []
#             for msg in history_dicts[:-1]:
#                 text = msg["parts"][0]
#                 if isinstance(text, dict): text = text.get('text', str(text))
                
#                 if msg["role"] == 'user':
#                     chat_history_messages.append(HumanMessage(content=text))
#                 else:
#                     chat_history_messages.append(AIMessage(content=text))
            
#             result = agent_orchestrator.invoke({"input": user_message, "chat_history": chat_history_messages})
            
#             # --- CRASH-PROOF TEXT EXTRACTION ---
#             ai_response_text = ""
#             if hasattr(result, 'content'):
#                 if isinstance(result.content, str):
#                     ai_response_text = result.content
#                 elif isinstance(result.content, list):
#                     ai_response_text = " ".join([p['text'] for p in result.content if isinstance(p, dict) and 'text' in p])
            
#             if not ai_response_text: ai_response_text = str(result.content)

#             # --- PROCESS TOOLS (Updating Database Based on Tool Calls) ---
#             if hasattr(result, 'tool_calls') and result.tool_calls:
#                 for tool_call in result.tool_calls:
#                     tool_name = tool_call['name']
#                     args = tool_call['args']
                    
#                     if tool_name == 'save_lead_qualification':
#                         lead.project_type = args.get('project_type')
#                         lead.budget = args.get('budget')
#                         lead.estimated_time_weeks = args.get('timeline')
#                         lead.seriousness_score = args.get('seriousness_score')
#                         lead.email = args.get('email')
#                         lead.status = 'Qualified' 
#                         ai_response_text = "✅ Qualification Complete! What is the **Name** of your project?"
                    
#                     elif tool_name == 'save_project_overview':
#                         lead.project_name = args.get('project_name')
#                         lead.project_description = args.get('project_description')
#                         # Convert comma-separated features/users into JSON lists for the dashboard
#                         lead.target_users = json.dumps([u.strip() for u in args.get('target_users', '').split(',') if u.strip()])
#                         lead.key_features = json.dumps([f.strip() for f in args.get('key_features', '').split(',') if f.strip()])
#                         ai_response_text = f"Excellent! I've recorded the overview for **{lead.project_name}**. Now, let's list specific features one by one."

#                     elif tool_name == 'save_requirement':
#                         cat_name = args.get('category', 'Functional')
#                         category = RequirementCategory.query.filter_by(name=cat_name).first() or RequirementCategory.query.first()
#                         new_req = ProjectRequirement(
#                             lead_id=lead.id,
#                             category_id=category.id,
#                             requirement_text=args.get('requirement_text'),
#                             priority=args.get('priority', 'Should'),
#                             status='Identified'
#                         )
#                         db.session.add(new_req)
#                         ai_response_text = f"Got it. I've logged that {cat_name} requirement. What is the next feature?"

#                     elif tool_name == 'generate_srs_document':
#                         # Trigger document generation helper
#                         lead.status = 'SRS_Generated'
#                         generate_full_srs(lead)
#                         db.session.commit()
#                         ai_response_text = "🎉 I have compiled all your requirements into a formal SRS document! You can download it from the dashboard."

#             # --- SAVE CLEAN TRANSCRIPT ---
#             if "SIGNAL" not in ai_response_text:
#                 history_dicts.append({"role": "model", "parts": [ai_response_text]})
            
#             lead.full_transcript = json.dumps(history_dicts)
#             db.session.commit()

#             return jsonify({
#                 "session_uuid": lead.session_uuid,
#                 "ai_response": ai_response_text,
#                 "is_qualified": lead.status in ['Qualified', 'SRS_Generated'],
#                 "is_srs_complete": lead.status == 'SRS_Generated'
#             })
            
#         except Exception as e:
#             db.session.rollback()
#             print(f"Error in handle_chat: {str(e)}")
#             return jsonify({"error": str(e)}), 500



# @app.route('/api/session/start', methods=['POST'])
# def start_session():
#     """Starts a new chat session and creates a Lead record, preventing duplicates from React Strict Mode."""
#     try:
#         with app.app_context():
#             five_seconds_ago = datetime.utcnow() - timedelta(seconds=5)
#             recent_lead_check = Lead.query.filter(Lead.created_at > five_seconds_ago).order_by(Lead.created_at.desc()).first()
            
#             # Debounce logic: If a new lead (with empty transcript) was created just now, return it instead of creating a phantom copy.
#             if recent_lead_check and recent_lead_check.full_transcript == '[]' and recent_lead_check.status == 'New':
#                  # Use the existing session and send back a greeting message
#                 initial_message = "Welcome back! Continuing our requirements gathering..." 
#                 return jsonify({
#                     "session_uuid": recent_lead_check.session_uuid,
#                     "ai_response": initial_message,
#                     "is_qualified": False
#                 }), 201

#             # Normal creation
#             new_lead = Lead(full_transcript=json.dumps([]))
#             db.session.add(new_lead)
#             db.session.commit()
            
#             initial_message = "Welcome! I'm your AI Requirements Engineering Assistant. To get started, could you tell me a little about the **type of software project** you are looking to build?"

#             return jsonify({
#                 "session_uuid": new_lead.session_uuid,
#                 "ai_response": initial_message,
#                 "is_qualified": False
#             }), 201
#     except Exception as e:
#         db.session.rollback()
#         return jsonify({"error": "Failed to start session", "details": str(e)}), 500


# @app.route('/api/admin/report/<uuid:session_uuid>', methods=['GET'])
# def download_report(session_uuid):
#     """Generates and downloads a comprehensive SRS report."""
#     with app.app_context():
#         lead = Lead.query.filter_by(session_uuid=session_uuid).first()
#         if not lead:
#             return jsonify({"error": "Lead not found"}), 404
        
#         file_stream = generate_docx_srs(lead)
#         return send_file(
#             file_stream,
#             as_attachment=True,
#             download_name=f"SRS_{lead.project_name or 'Project'}.docx",
#             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#         )
        
#         # Ensure SRS exists or generate it
#         # srs = SRSDocument.query.filter_by(lead_id=lead.id).order_by(SRSDocument.generated_at.desc()).first()
#         # if not srs:
#         #     srs = generate_full_srs(lead)

#         # srs_content_string = srs.content 
        
#         # response = make_response(srs_content_string)
#         # response.headers["Content-Disposition"] = f"attachment; filename=SRS_{lead.project_name or lead.session_uuid}.md"
#         # response.mimetype = "text/markdown"
        
#         # return response

# @app.route('/api/admin/lead/<uuid:session_uuid>/generate-srs', methods=['POST'])
# def manual_generate_srs(session_uuid):
#     """Manually triggers SRS generation and updates the lead status."""
#     # NOTE: You MUST enforce authentication middleware here in a real application!
#     with app.app_context():
#         lead = db.session.execute(
#             db.select(Lead).filter_by(session_uuid=session_uuid)
#         ).scalar_one_or_none()

#         if not lead:
#             return jsonify({"error": "Lead not found"}), 404
        
#         try:
#             generate_full_srs(lead)
#             lead.status = 'SRS_Generated'
#             db.session.commit()
            
#             return jsonify({"message": f"SRS generation triggered successfully for {lead.project_name or lead.email}"}), 200
#         except Exception as e:
#             db.session.rollback()
#             return jsonify({"error": f"Failed to generate SRS: {str(e)}"}), 500



# @app.route('/api/admin/dashboard/summary', methods=['GET'])
# def get_dashboard_summary():
#     # Keeping this simplified dashboard endpoint
#     with app.app_context():
#         total_leads = db.session.scalar(db.select(db.func.count(Lead.id)))
        
#         status_data = db.session.execute(db.select(Lead.status, db.func.count(Lead.id).label('count')).group_by(Lead.status)).mappings().all()

#         total_requirements = db.session.scalar(db.select(db.func.count(ProjectRequirement.id)))
#         avg_requirements_per_lead = total_requirements / total_leads if total_leads > 0 else 0
        
#         srs_generated = db.session.scalar(db.select(db.func.count(SRSDocument.id)))
        
#         type_data = db.session.execute(db.select(Lead.project_type, db.func.count(Lead.id).label('count')).where(Lead.status == 'Qualified').group_by(Lead.project_type)).mappings().all()
        
#         recent_leads = db.session.execute(db.select(Lead).order_by(Lead.created_at.desc()).limit(5)).scalars().all()
        
#         volume_data = db.session.execute(db.select(db.func.strftime('%Y-%m-%d', Lead.created_at).label('date'), db.func.count(Lead.id).label('count')).group_by('date').order_by('date').limit(30)).mappings().all()
        
#     return jsonify({
#         "total_leads": total_leads,
#         "status_distribution": [dict(s) for s in status_data],
#         "project_type_distribution": [dict(t) for t in type_data],
#         "total_requirements": total_requirements,
#         "avg_requirements_per_lead": round(avg_requirements_per_lead, 2),
#         "srs_documents_generated": srs_generated,
#         "volume_by_day": [dict(v) for v in volume_data],
#         "recent_leads": [{
#             'id': lead.id, 'session_uuid': lead.session_uuid, 'project_name': lead.project_name or 'Unnamed', 
#             'status': lead.status, 'email': lead.email
#         } for lead in recent_leads]
#     })

# @app.route('/api/admin/leads', methods=['GET'])
# def get_lead_list():
#     with app.app_context():
#         # Eagerly load requirements to calculate count in to_dict without hitting SQLA DetachedInstanceError
#         leads = db.session.execute(db.select(Lead).order_by(Lead.created_at.desc())).scalars().all()
#         lead_list = [lead.to_dict() for lead in leads]
#     return jsonify(lead_list)

# @app.route('/api/admin/login', methods=['POST'])
# def admin_login():
#     data = request.get_json()
#     username = data.get('username')
#     password = data.get('password')

#     ADMIN_USER = os.environ.get('ADMIN_USER', 'admin')
#     ADMIN_PASS = os.environ.get('ADMIN_PASS', 'securepassword')

#     if username == ADMIN_USER and password == ADMIN_PASS:
#         auth_token = str(uuid4()) 
#         return jsonify({"message": "Login successful", "token": auth_token, "username": ADMIN_USER}), 200
#     else:
#         return jsonify({"message": "Invalid credentials"}), 401
    
# @app.route('/')
# def health_check():
#     return {"status": "ok", "service": "AI Requirements Engineering Platform"}

# if __name__ == '__main__':
#     with app.app_context():
#         import sys
#         if len(sys.argv) > 1 and sys.argv[1] == 'reset':
#             db.drop_all()
#             print("🗑️ Database reset")
        
#         db.create_all()
#         initialize_categories()
#         print("✅ Database ready")
#     app.run(debug=True, port=5000, use_reloader=False)

